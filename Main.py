import streamlit as st
import os
import PyPDF2
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_core.prompts import ChatPromptTemplate
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_groq import ChatGroq
from langchain_google_genai import ChatGoogleGenerativeAI
from fpdf import FPDF
import re
import random
import time
import concurrent.futures
import pytesseract
from pdf2image import convert_from_bytes
from PIL import Image
import io
import requests
from datetime import datetime
import tempfile
import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd

# ─── Constants ───────────────────────────────────────────────────────────────

EDUCATION_LEVELS = ["School", "College"]

SCHOOL_BOARDS = ["CBSE", "ICSE", "State Board"]
COLLEGE_BOARDS = ["University/College", "UGC", "AICTE", "Autonomous"]

SCHOOL_CLASSES = list(range(1, 13))
COLLEGE_YEARS = ["1st Year", "2nd Year", "3rd Year", "4th Year", "Post Graduate"]

SCHOOL_SUBJECTS = [
    "Mathematics", "Science", "Social Studies", "English", "Hindi",
    "Physics", "Chemistry", "Biology", "Computer Science", "Other"
]

COLLEGE_SUBJECTS = [
    # Engineering & CS
    "Data Structures & Algorithms", "Database Management Systems",
    "Operating Systems", "Computer Networks", "Machine Learning",
    "Artificial Intelligence", "Software Engineering", "Web Development",
    "Cyber Security", "Cloud Computing",
    # Science
    "Physics", "Chemistry", "Organic Chemistry", "Inorganic Chemistry",
    "Biology", "Biotechnology", "Microbiology", "Biochemistry",
    "Mathematics", "Statistics", "Discrete Mathematics",
    # Commerce & Management
    "Accounting", "Business Studies", "Economics", "Marketing",
    "Financial Management", "Human Resource Management",
    "Business Law", "Entrepreneurship",
    # Arts & Humanities
    "English Literature", "Psychology", "Sociology", "Political Science",
    "History", "Philosophy", "Journalism",
    # Engineering Core
    "Engineering Mechanics", "Thermodynamics", "Fluid Mechanics",
    "Electrical Engineering", "Electronics", "Signals & Systems",
    "Control Systems", "Digital Logic Design",
    "Mechanical Engineering", "Civil Engineering",
    # Medical
    "Anatomy", "Physiology", "Pharmacology", "Pathology",
    # Other
    "Other"
]

SCHOOL_QUESTION_TYPES = [
    "multiple-choice", "short-answer", "long-answer",
    "true-false", "fill-in-the-blanks", "match-the-following"
]

COLLEGE_QUESTION_TYPES = [
    "multiple-choice", "short-answer", "long-answer",
    "true-false", "fill-in-the-blanks", "match-the-following",
    "case-study", "numerical-problem", "diagram-based",
    "assertion-reason", "conceptual-analysis"
]

COMPLEXITY_LEVELS = {
    "School": ["basic", "intermediate", "advanced"],
    "College": ["foundational", "intermediate", "advanced", "research-level"]
}


# ─── Session State ───────────────────────────────────────────────────────────

def init_session_state():
    defaults = {
        'vectors': None,
        'pdf_ready': False,
        'answers_generated': False,
        'question_settings': None,
        'ocr_enabled': True,
        'raw_text': None,
        'ai_answer': None,
        'answers': None,
        'question_bank': [],
        'generation_history': [],
        'education_level': 'School'
    }

    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value


# ─── Helpers ─────────────────────────────────────────────────────────────────

def get_api_key(key_name):
    if key_name in st.secrets:
        return st.secrets[key_name]
    elif key_name in os.environ:
        return os.environ[key_name]
    else:
        st.error(f"Missing API key: {key_name}. Please add it to your secrets.")
        return None


@st.cache_resource
def load_embeddings():
    """Load HuggingFace embeddings model (cached for performance)"""
    return HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-mpnet-base-v2"
    )


def estimate_tokens(text):
    words = text.split()
    return int(len(words) * 1.3)


def get_level_label(class_level):
    """Return a human-readable label for the class/year level."""
    if isinstance(class_level, int):
        return f"Class {class_level}"
    return str(class_level)


# ─── Validation ──────────────────────────────────────────────────────────────

def validate_inputs(board, class_level, subject, num_questions, question_type, education_level="School"):
    errors = []

    valid_boards = SCHOOL_BOARDS + COLLEGE_BOARDS
    if not board or board not in valid_boards:
        errors.append("⚠️ Invalid educational board selected")

    if education_level == "School":
        if isinstance(class_level, int):
            if not (1 <= class_level <= 12):
                errors.append("⚠️ Class level must be between 1 and 12")
        else:
            errors.append("⚠️ Invalid class level for school")
    else:
        if class_level not in COLLEGE_YEARS:
            errors.append("⚠️ Invalid year selected for college")

    valid_subjects = SCHOOL_SUBJECTS + COLLEGE_SUBJECTS
    if not subject or subject not in valid_subjects:
        errors.append("⚠️ Invalid subject selected")

    if not (5 <= num_questions <= 25):
        errors.append("⚠️ Number of questions must be between 5 and 25")

    valid_question_types = SCHOOL_QUESTION_TYPES + COLLEGE_QUESTION_TYPES
    if not question_type or question_type not in valid_question_types:
        errors.append("⚠️ Invalid question type selected")

    return len(errors) == 0, errors


# ─── PDF Processing ─────────────────────────────────────────────────────────

def is_pdf_scanned(pdf_bytes):
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
        first_page = pdf_reader.pages[0]
        text = first_page.extract_text()
        if len(text.strip()) < 100:
            return True
        return False
    except Exception:
        return True


def extract_text_from_image(image):
    try:
        text = pytesseract.image_to_string(image)
        return text
    except Exception as e:
        st.error(f"OCR Error: {e}")
        return ""


def process_scanned_pdf(pdf_bytes):
    try:
        images = convert_from_bytes(pdf_bytes)
        with st.spinner("Performing OCR on scanned PDF..."):
            with concurrent.futures.ThreadPoolExecutor() as executor:
                text_chunks = list(executor.map(extract_text_from_image, images))
        combined_text = "\n\n".join([chunk for chunk in text_chunks if chunk])
        return combined_text
    except Exception as e:
        st.error(f"Failed to process scanned PDF: {e}")
        return ""


def get_pdf_text(upload_pdf):
    if upload_pdf is None:
        st.error("No PDF file uploaded.")
        return None

    try:
        pdf_bytes = upload_pdf.getvalue()

        if len(pdf_bytes) > 50 * 1024 * 1024:
            st.error("PDF file is too large. Please upload a file smaller than 50MB.")
            return None

        if is_pdf_scanned(pdf_bytes):
            st.info("Detected a scanned/image-based PDF. Using OCR to extract text...")
            return process_scanned_pdf(pdf_bytes)

        pdf_reader = PyPDF2.PdfReader(upload_pdf)
        extracted_text = ""
        for page in pdf_reader.pages:
            text = page.extract_text()
            if text:
                extracted_text += text + "\n\n"

        return extracted_text
    except Exception as e:
        st.error(f"Failed to extract text from PDF: {e}")
        return None


# ─── Chunking & Vector DB ───────────────────────────────────────────────────

def split_data_into_chunks(raw_text):
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        separators=["\n\n", "\n", " ", ""],
        length_function=len
    )
    chunks = text_splitter.split_text(raw_text)
    return chunks


def store_chunks_vectorDB(chunks, embedding):
    batch_size = 100
    vector_db = None
    progress_bar = st.progress(0)
    status_text = st.empty()
    total_batches = (len(chunks) + batch_size - 1) // batch_size

    for i in range(0, len(chunks), batch_size):
        batch_chunks = chunks[i:i + batch_size]
        batch_num = i // batch_size
        progress_bar.progress((batch_num + 1) / total_batches)
        status_text.text(f"Processing batch {batch_num + 1} of {total_batches}...")

        try:
            if vector_db is None:
                vector_db = FAISS.from_texts(texts=batch_chunks, embedding=embedding)
            else:
                vector_db.add_texts(texts=batch_chunks)
            time.sleep(0.1)
        except Exception as e:
            st.error(f"Error processing batch {batch_num + 1}: {e}")
            raise

    progress_bar.empty()
    status_text.empty()
    return vector_db


# ─── Export Utilities ────────────────────────────────────────────────────────

def parse_questions_to_list(content):
    lines = content.split('\n')
    questions_list = []
    current_question = ""
    current_number = 0

    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line[0].isdigit() and '.' in line[:3]:
            if current_question:
                questions_list.append({
                    'number': current_number,
                    'text': current_question.strip()
                })
            current_number += 1
            current_question = line
        else:
            current_question += " " + line

    if current_question:
        questions_list.append({
            'number': current_number,
            'text': current_question.strip()
        })

    return questions_list


def create_word_doc(content, board, class_level, subject, question_type, is_answer_key=False):
    doc = Document()
    level_label = get_level_label(class_level)

    title = doc.add_heading(f"{board} {level_label} - {subject}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle_text = f"{question_type.title()} {'Answer Key' if is_answer_key else 'Questions'}"
    subtitle = doc.add_heading(subtitle_text, level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f"Generated on: {datetime.now().strftime('%d-%m-%Y')}")
    date_run.italic = True
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()

    for line in content.split('\n'):
        if line.strip():
            if line.strip().startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')) or \
               line.strip().startswith(('Question', 'Q.', 'I.', 'II.', 'Case Study', 'Scenario')):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.size = Pt(12)
            elif line.strip().startswith(('A)', 'B)', 'C)', 'D)', 'a)', 'b)', 'c)', 'd)')):
                p = doc.add_paragraph(line, style='List Bullet')
            else:
                doc.add_paragraph(line)

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()
    return temp_file.name


def create_excel_file(content, board, class_level, subject, question_type, is_answer_key=False):
    questions_list = parse_questions_to_list(content)
    level_label = get_level_label(class_level)

    data = []
    for q in questions_list:
        data.append({
            'No.': q['number'],
            'Question/Answer': q['text'],
            'Type': question_type,
            'Category': 'Answer' if is_answer_key else 'Question'
        })

    df = pd.DataFrame(data)

    metadata = {
        'Board': [board],
        'Level': [level_label],
        'Subject': [subject],
        'Question Type': [question_type],
        'Generated On': [datetime.now().strftime('%d-%m-%Y %H:%M:%S')],
        'Total Items': [len(data)]
    }
    df_metadata = pd.DataFrame(metadata)

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx')
    with pd.ExcelWriter(temp_file.name, engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name='Questions', index=False)
        df_metadata.to_excel(writer, sheet_name='Metadata', index=False)
    temp_file.close()
    return temp_file.name


def create_json_export(content, settings):
    questions_list = parse_questions_to_list(content)
    export_data = {
        'metadata': {
            'board': settings.get('board', 'Unknown'),
            'level': settings.get('class', 'Unknown'),
            'education_level': settings.get('education_level', 'Unknown'),
            'subject': settings.get('subject', 'Unknown'),
            'question_type': settings.get('type', 'Unknown'),
            'complexity': settings.get('complexity', 'Unknown'),
            'generated_at': datetime.now().isoformat(),
            'total_questions': len(questions_list)
        },
        'questions': questions_list
    }
    return json.dumps(export_data, indent=2, ensure_ascii=False)


def save_to_question_bank(questions, answers, settings):
    if 'question_bank' not in st.session_state:
        st.session_state.question_bank = []

    entry = {
        'id': len(st.session_state.question_bank) + 1,
        'timestamp': datetime.now().isoformat(),
        'settings': settings,
        'questions': questions,
        'answers': answers,
        'tags': [settings['board'], settings['subject'], settings['type'],
                 settings.get('education_level', 'School')]
    }
    st.session_state.question_bank.append(entry)

    if 'generation_history' not in st.session_state:
        st.session_state.generation_history = []

    st.session_state.generation_history.append({
        'timestamp': datetime.now(),
        'settings': settings,
        'question_count': len(parse_questions_to_list(questions)),
        'preview': questions[:100] + "..."
    })


def create_temp_pdf(content, board, class_level, subject, question_type, is_answer_key=False):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    level_label = get_level_label(class_level)

    pdf.set_font("Arial", size=12)
    pdf.set_font("Arial", style="B", size=16)

    def clean_text(text):
        if isinstance(text, str):
            return text.encode('latin-1', 'replace').decode('latin-1')
        return str(text)

    header_text = clean_text(
        f"{board} {level_label} - {subject} {question_type.title()} "
        f"{'Answer Key' if is_answer_key else 'Questions'}"
    )
    pdf.cell(200, 10, txt=header_text, ln=True, align="C")
    pdf.ln(5)

    pdf.set_font("Arial", style="I", size=10)
    date_text = clean_text(f"Generated on: {datetime.now().strftime('%d-%m-%Y')}")
    pdf.cell(200, 10, txt=date_text, ln=True, align="R")
    pdf.ln(10)

    pdf.set_font("Arial", size=12)
    for line in content.split('\n'):
        clean_line = clean_text(line)
        if clean_line.strip().startswith(("Question", "Q.", "1.", "2.", "I.", "II.", "Case Study", "Scenario")):
            pdf.set_font("Arial", style="B", size=12)
            pdf.cell(0, 7, txt=clean_line, ln=True, align="L")
            pdf.ln(2)
        elif clean_line.startswith(("A.", "B.", "C.", "D.", "a)", "b)", "c)", "d)", "(a)", "(b)")):
            pdf.set_font("Arial", size=12)
            pdf.cell(10, 5, txt="", ln=0)
            pdf.cell(0, 5, txt=clean_line, ln=True)
        else:
            pdf.set_font("Arial", size=12)
            pdf.multi_cell(0, 5, txt=clean_line)
            pdf.ln(1)

    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
        temp_path = temp_file.name
    pdf.output(temp_path)
    return temp_path


# ─── Prompt Builders ─────────────────────────────────────────────────────────

def create_diverse_question_prompt(question_type, board, class_level, subject, seed=None, education_level="School"):
    if seed is None:
        seed = random.randint(1, 1000)

    level_label = get_level_label(class_level)

    # Determine depth / style based on education level
    if education_level == "College":
        level_description = f"college/university {level_label} level"
        depth = "in-depth, analytical, application-oriented, and aligned with higher-education standards"
        exam_style = "university semester exams, competitive exams, and professional certification tests"
    else:
        level_description = f"{board} {level_label}"
        depth = "age-appropriate, curriculum-aligned, and suitable for school board exams"
        exam_style = f"{board} board examinations"

    approaches = [
        "analytical thinking",
        "critical reasoning",
        "real-world application",
        "conceptual understanding",
        "creative problem-solving"
    ]

    if subject in ["Mathematics", "Discrete Mathematics", "Statistics"]:
        focus_areas = random.sample([
            "algebraic concepts", "geometric principles", "statistical analysis",
            "numerical reasoning", "mathematical proofs", "applied mathematics",
            "probability theory", "linear algebra", "calculus concepts"
        ], k=2)
    elif subject in ["Physics", "Chemistry", "Biology", "Science",
                      "Organic Chemistry", "Inorganic Chemistry", "Biotechnology",
                      "Microbiology", "Biochemistry"]:
        focus_areas = random.sample([
            "scientific principles", "experimental design", "natural phenomena",
            "scientific theories", "laboratory techniques", "scientific discoveries",
            "research methodology", "applied sciences", "interdisciplinary concepts"
        ], k=2)
    elif subject in ["Data Structures & Algorithms", "Database Management Systems",
                      "Operating Systems", "Computer Networks", "Machine Learning",
                      "Artificial Intelligence", "Software Engineering", "Web Development",
                      "Cyber Security", "Cloud Computing", "Computer Science",
                      "Digital Logic Design", "Signals & Systems", "Control Systems"]:
        focus_areas = random.sample([
            "algorithm analysis", "system design", "data modeling",
            "computational complexity", "practical implementation",
            "real-world applications", "optimization techniques",
            "architectural patterns", "security considerations"
        ], k=2)
    elif subject in ["Accounting", "Business Studies", "Economics", "Marketing",
                      "Financial Management", "Human Resource Management",
                      "Business Law", "Entrepreneurship"]:
        focus_areas = random.sample([
            "financial analysis", "market dynamics", "regulatory frameworks",
            "strategic planning", "case-based reasoning", "economic models",
            "business ethics", "quantitative methods", "management principles"
        ], k=2)
    elif subject in ["Anatomy", "Physiology", "Pharmacology", "Pathology"]:
        focus_areas = random.sample([
            "clinical reasoning", "diagnostic thinking", "anatomical relationships",
            "physiological mechanisms", "drug interactions", "disease pathways",
            "patient case analysis", "treatment protocols"
        ], k=2)
    else:
        focus_areas = random.sample([
            "key concepts", "important theories", "significant events",
            "critical analysis", "practical applications", "foundational principles",
            "comparative analysis", "contemporary relevance"
        ], k=2)

    cognitive_levels = random.sample([
        "knowledge recall", "comprehension", "application",
        "analysis", "synthesis", "evaluation"
    ], k=3)

    base_prompt = f"""
    You are an expert {level_description} {subject} teacher with deep knowledge of {exam_style}.

    Generate ONLY {question_type} questions based on the context provided.
    The questions must be {depth}.

    To ensure DIVERSE and HIGH-QUALITY questions:
    - Focus on {focus_areas[0]} and {focus_areas[1]}
    - Target cognitive levels of {cognitive_levels[0]}, {cognitive_levels[1]}, and {cognitive_levels[2]}
    - Use an approach emphasizing {random.choice(approaches)}
    - Ensure questions are varied in difficulty and style
    - Use creativity seed #{seed} to make questions unique and non-repetitive

    IMPORTANT:
    1. Create ONLY {question_type} questions - do not mix question types
    2. Do not include any explanations, notes, or additional text
    3. Start directly with "1." - no preamble
    4. Number your questions sequentially

    Context: {{context}}
    Complexity Level: {{complexity_level}}
    Number of Questions: {{num_questions}}
    """

    # ── Format instructions per question type ────────────────────────────────

    if question_type == "multiple-choice":
        base_prompt += """
        FORMAT STRICTLY AS:
        1. [Question text]?
           A) [Option A]
           B) [Option B]
           C) [Option C]
           D) [Option D]

        2. [Next question]?
           A) [Option A]
           B) [Option B]
           C) [Option C]
           D) [Option D]

        and so on.
        """

    elif question_type == "short-answer":
        base_prompt += """
        FORMAT STRICTLY AS:
        1. [Question text]? (Answer in 30-50 words)

        2. [Next question]? (Answer in 30-50 words)

        and so on.
        """

    elif question_type == "long-answer":
        word_range = "150-250 words" if education_level == "College" else "100-150 words"
        base_prompt += f"""
        FORMAT STRICTLY AS:
        1. [Question text]? (Answer in {word_range})

        2. [Next question]? (Answer in {word_range})

        and so on.
        """

    elif question_type == "true-false":
        base_prompt += """
        FORMAT STRICTLY AS:
        1. [Statement]. (True/False)

        2. [Statement]. (True/False)

        and so on.
        """

    elif question_type == "fill-in-the-blanks":
        base_prompt += """
        FORMAT STRICTLY AS:
        1. [Sentence with __________ as blank]

        2. [Sentence with __________ as blank]

        and so on.
        """

    elif question_type == "match-the-following":
        base_prompt += """
        FORMAT STRICTLY AS:
        1. Match the following:
           Column A              Column B
           a) [Item 1]          i) [Match A]
           b) [Item 2]          ii) [Match B]
           c) [Item 3]          iii) [Match C]
           d) [Item 4]          iv) [Match D]

        2. Match the following:
           Column A              Column B
           a) [Item 1]          i) [Match A]
           b) [Item 2]          ii) [Match B]
           c) [Item 3]          iii) [Match C]
           d) [Item 4]          iv) [Match D]

        and so on.
        """

    elif question_type == "case-study":
        base_prompt += """
        FORMAT STRICTLY AS:
        1. Case Study:
           [A detailed scenario/case description in 80-120 words relevant to the subject]

           Based on the above case study, answer the following:
           a) [Sub-question 1]
           b) [Sub-question 2]
           c) [Sub-question 3]

        2. Case Study:
           [A detailed scenario/case description in 80-120 words relevant to the subject]

           Based on the above case study, answer the following:
           a) [Sub-question 1]
           b) [Sub-question 2]
           c) [Sub-question 3]

        and so on.
        """

    elif question_type == "numerical-problem":
        base_prompt += """
        FORMAT STRICTLY AS:
        1. [Numerical problem statement with all given data and units clearly mentioned]
           (Show complete working/steps in your answer)

        2. [Numerical problem statement with all given data and units clearly mentioned]
           (Show complete working/steps in your answer)

        and so on.
        """

    elif question_type == "diagram-based":
        base_prompt += """
        FORMAT STRICTLY AS:
        1. [Question that requires drawing, labelling, or interpreting a diagram]
           Instruction: [Specify what kind of diagram is expected, e.g., "Draw and label...", "Sketch the circuit for...", "Illustrate the process of..."]

        2. [Question that requires drawing, labelling, or interpreting a diagram]
           Instruction: [Specify what kind of diagram is expected]

        and so on.
        """

    elif question_type == "assertion-reason":
        base_prompt += """
        FORMAT STRICTLY AS:
        1. Assertion (A): [A factual statement]
           Reason (R): [An explanation or justification]

           Choose the correct option:
           A) Both A and R are true, and R is the correct explanation of A
           B) Both A and R are true, but R is NOT the correct explanation of A
           C) A is true but R is false
           D) A is false but R is true

        2. Assertion (A): [A factual statement]
           Reason (R): [An explanation or justification]

           Choose the correct option:
           A) Both A and R are true, and R is the correct explanation of A
           B) Both A and R are true, but R is NOT the correct explanation of A
           C) A is true but R is false
           D) A is false but R is true

        and so on.
        """

    elif question_type == "conceptual-analysis":
        base_prompt += """
        FORMAT STRICTLY AS:
        1. [In-depth analytical question requiring comparison, evaluation, or critical thinking about concepts]
           (Answer should demonstrate deep understanding with examples in 150-250 words)

        2. [In-depth analytical question requiring comparison, evaluation, or critical thinking about concepts]
           (Answer should demonstrate deep understanding with examples in 150-250 words)

        and so on.
        """

    return base_prompt


def get_answer_prompt(question_type, board, class_level, subject, education_level="School"):
    level_label = get_level_label(class_level)

    if education_level == "College":
        level_description = f"college/university {level_label} level"
        detail_note = "Provide thorough, academically rigorous answers suitable for university-level evaluation."
    else:
        level_description = f"{board} {level_label}"
        detail_note = "Provide clear, accurate, and educational answers suitable for school board evaluation."

    base_prompt = f"""
    You are an expert {level_description} {subject} teacher.

    Below are {question_type} questions. Provide correct answers for each question.
    {detail_note}

    Questions:
    {{questions}}

    IMPORTANT:
    1. Provide answers in same numbered order as questions
    2. Start directly with "1." - no preamble
    3. Be clear, accurate, and educational in your answers
    """

    if question_type == "multiple-choice":
        base_prompt += """
        FORMAT STRICTLY AS:
        1. [Correct Option Letter]: [Brief explanation why this is correct]

        2. [Correct Option Letter]: [Brief explanation why this is correct]

        and so on.
        """

    elif question_type in ("short-answer", "long-answer", "conceptual-analysis"):
        base_prompt += """
        FORMAT STRICTLY AS:
        1. [Complete answer that directly addresses the question]

        2. [Complete answer that directly addresses the question]

        and so on.
        """

    elif question_type == "true-false":
        base_prompt += """
        FORMAT STRICTLY AS:
        1. [TRUE or FALSE]: [Brief explanation supporting this answer]

        2. [TRUE or FALSE]: [Brief explanation supporting this answer]

        and so on.
        """

    elif question_type == "fill-in-the-blanks":
        base_prompt += """
        FORMAT STRICTLY AS:
        1. [Word or phrase that fills the blank]: [Brief explanation if needed]

        2. [Word or phrase that fills the blank]: [Brief explanation if needed]

        and so on.
        """

    elif question_type == "match-the-following":
        base_prompt += """
        FORMAT STRICTLY AS:
        1. a-[correct number], b-[correct number], c-[correct number], d-[correct number]: [Brief explanation of relationships]

        2. a-[correct number], b-[correct number], c-[correct number], d-[correct number]: [Brief explanation of relationships]

        and so on.
        """

    elif question_type == "case-study":
        base_prompt += """
        FORMAT STRICTLY AS:
        1. Case Study Answers:
           a) [Detailed answer to sub-question 1]
           b) [Detailed answer to sub-question 2]
           c) [Detailed answer to sub-question 3]

        2. Case Study Answers:
           a) [Detailed answer to sub-question 1]
           b) [Detailed answer to sub-question 2]
           c) [Detailed answer to sub-question 3]

        and so on.
        """

    elif question_type == "numerical-problem":
        base_prompt += """
        FORMAT STRICTLY AS:
        1. Given: [List given values]
           To Find: [What needs to be calculated]
           Solution:
           [Step-by-step working with formulas and calculations]
           Answer: [Final answer with units]

        2. Given: [List given values]
           To Find: [What needs to be calculated]
           Solution:
           [Step-by-step working with formulas and calculations]
           Answer: [Final answer with units]

        and so on.
        """

    elif question_type == "diagram-based":
        base_prompt += """
        FORMAT STRICTLY AS:
        1. [Detailed text description of the diagram with all labels and parts identified]
           Key Points: [Important features to include in the diagram]

        2. [Detailed text description of the diagram with all labels and parts identified]
           Key Points: [Important features to include in the diagram]

        and so on.
        """

    elif question_type == "assertion-reason":
        base_prompt += """
        FORMAT STRICTLY AS:
        1. [Correct Option Letter]: [Explanation of why the assertion is true/false and how the reason relates to it]

        2. [Correct Option Letter]: [Explanation of why the assertion is true/false and how the reason relates to it]

        and so on.
        """

    return base_prompt


# ─── Format Checking ─────────────────────────────────────────────────────────

def check_question_format(text, question_type):
    lower_text = text.lower()

    checks = {
        "multiple-choice": lambda t: any(x in t for x in ["a)", "a.", "a )"]),
        "short-answer": lambda t: "words" in t or "answer" in t,
        "long-answer": lambda t: "words" in t or "100" in t or "150" in t,
        "true-false": lambda t: "true/false" in t or "true or false" in t,
        "fill-in-the-blanks": lambda t: "____" in t or "..." in t or "_____" in t,
        "match-the-following": lambda t: "column" in t,
        "case-study": lambda t: "case study" in t or "scenario" in t or "based on" in t,
        "numerical-problem": lambda t: "calculate" in t or "find" in t or "determine" in t or "given" in t,
        "diagram-based": lambda t: "draw" in t or "label" in t or "diagram" in t or "sketch" in t or "illustrate" in t,
        "assertion-reason": lambda t: "assertion" in t or "reason" in t,
        "conceptual-analysis": lambda t: "analyze" in t or "compare" in t or "evaluate" in t or "discuss" in t,
    }

    checker = checks.get(question_type, lambda t: True)
    return checker(lower_text)


# ─── Post-processing ────────────────────────────────────────────────────────

def post_process_output(text):
    patterns = [
        r'\n\s*1\.\s',
        r'\n\s*Question\s*1[\s\:\.]',
        r'\n\s*Q\.?\s*1[\s\:\.]',
    ]
    earliest_pos = len(text)
    for pattern in patterns:
        matches = re.search(pattern, text)
        if matches and matches.start() < earliest_pos:
            earliest_pos = matches.start()
    if earliest_pos < len(text):
        return text[earliest_pos + 1:].strip()

    question_markers = re.search(
        r'\b(what|which|how|why|where|when|describe|explain|calculate|determine|identify|list|case study|assertion|given)\b',
        text.lower()
    )
    if question_markers:
        return text[question_markers.start():].strip()

    return text


# ─── Generation & Validation ────────────────────────────────────────────────

def generate_questions(inputs, max_retries=2):
    content = inputs['content']
    question_type = inputs['question_types']
    complexity_level = inputs['complexity_level']
    num_questions = inputs['num_questions']
    board = inputs.get('board', 'CBSE')
    class_level = inputs.get('class_level', '10')
    subject = inputs.get('subject', 'General')
    education_level = inputs.get('education_level', 'School')

    is_valid, errors = validate_inputs(board, class_level, subject, num_questions, question_type, education_level)
    if not is_valid:
        for error in errors:
            st.error(error)
        return "Please correct the errors above and try again."

    level_label = get_level_label(class_level)

    st.info(f"""
    Generating {num_questions} diverse {question_type} questions
    for {board} {level_label} {subject} ({education_level}) at {complexity_level} complexity
    """)

    st.session_state.question_settings = {
        "type": question_type,
        "board": board,
        "class": class_level,
        "subject": subject,
        "complexity": complexity_level,
        "count": num_questions,
        "education_level": education_level
    }

    groq_api_key = get_api_key("GROQ_API_KEY")
    if not groq_api_key:
        st.error("GROQ API key is missing. Please check your secrets.")
        return "Error: Missing API key. Contact the administrator."

    llm = ChatGroq(api_key=groq_api_key, model='llama-3.3-70b-versatile')

    seed = random.randint(1, 1000)

    prompt_template = create_diverse_question_prompt(
        question_type, board, class_level, subject,
        seed=seed, education_level=education_level
    )
    prompt = ChatPromptTemplate.from_template(prompt_template)

    if st.session_state.vectors:
        search_query = f"{board} {level_label} {subject} {question_type} questions {complexity_level} complexity"

        with st.spinner('Retrieving relevant content from document...'):
            k_chunks = 5
            relevant_docs = st.session_state.vectors.similarity_search(search_query, k=k_chunks)
            relevant_content = "\n\n".join([doc.page_content for doc in relevant_docs])

            estimated_tokens = estimate_tokens(relevant_content)

            while estimated_tokens > 2000 and k_chunks > 1:
                k_chunks -= 1
                relevant_docs = st.session_state.vectors.similarity_search(search_query, k=k_chunks)
                relevant_content = "\n\n".join([doc.page_content for doc in relevant_docs])
                estimated_tokens = estimate_tokens(relevant_content)

            st.session_state.content_tokens = estimated_tokens

            if estimated_tokens > 2000:
                words = relevant_content.split()
                relevant_content = " ".join(words[:1500])

            content = relevant_content
    else:
        words = content.split()
        if len(words) > 1500:
            content = " ".join(words[:1500])

    for attempt in range(max_retries + 1):
        try:
            if attempt > 0:
                st.info(f"Retry attempt {attempt} with a different approach...")

            with st.spinner(f'AI is generating {question_type} questions (attempt {attempt + 1})...'):
                formatted_prompt = prompt.format(
                    context=content,
                    complexity_level=complexity_level,
                    num_questions=num_questions
                )

                prompt_tokens = estimate_tokens(formatted_prompt)
                if prompt_tokens > 5000:
                    st.warning(f"Prompt too large (~{prompt_tokens} tokens). Reducing content size...")
                    words = content.split()
                    reduced_words = int(len(words) * (4000 / prompt_tokens))
                    content = " ".join(words[:reduced_words])
                    formatted_prompt = prompt.format(
                        context=content,
                        complexity_level=complexity_level,
                        num_questions=num_questions
                    )

                start_time = time.time()
                response = llm.invoke(formatted_prompt)
                generation_time = time.time() - start_time

                st.session_state.last_generation_time = generation_time

            ai_answer = response.content if hasattr(response, 'content') else str(response)
            cleaned_answer = post_process_output(ai_answer)

            if check_question_format(cleaned_answer, question_type):
                st.success(f"Successfully generated {question_type} questions in {generation_time:.1f} seconds!")
                st.session_state.generated_questions = cleaned_answer
                return cleaned_answer

            if attempt == max_retries:
                st.warning("Output format needs adjustment. Reformatting...")
                fix_prompt = f"""
                I need {num_questions} {question_type} questions for {board} {level_label} {subject}.

                The content I have doesn't match the {question_type} format.

                Please reformat the following content into proper {question_type} questions:

                {cleaned_answer}

                Make sure to follow the proper format for {question_type} questions as described earlier.
                """
                with st.spinner("Reformatting questions..."):
                    fix_response = llm.invoke(fix_prompt)
                    cleaned_answer = post_process_output(fix_response.content)

                st.session_state.generated_questions = cleaned_answer
                return cleaned_answer

            seed = random.randint(1, 1000)
            prompt_template = create_diverse_question_prompt(
                question_type, board, class_level, subject,
                seed=seed, education_level=education_level
            )
            prompt = ChatPromptTemplate.from_template(prompt_template)

        except Exception as e:
            st.error(f"Error during generation (attempt {attempt + 1}): {e}")

            error_str = str(e)
            if "413" in error_str or "token" in error_str.lower() or "too large" in error_str.lower():
                words = content.split()
                content = " ".join(words[:len(words) // 2])
                st.warning("Request was too large. Reducing content size and trying again...")

            if attempt == max_retries:
                fallback_content = subject + " " + board + " curriculum highlights"
                st.warning("Using minimal content as fallback...")
                try:
                    fallback_prompt = prompt.format(
                        context=fallback_content,
                        complexity_level=complexity_level,
                        num_questions=num_questions
                    )
                    response = llm.invoke(fallback_prompt)
                    fallback_answer = post_process_output(response.content)
                    st.session_state.generated_questions = fallback_answer
                    return fallback_answer
                except Exception:
                    st.error("Maximum retries reached. Please try again with less content.")
                    return "Failed to generate questions. Please try again with a smaller document or fewer questions."

    return "Failed to generate questions after multiple attempts."


def generate_answers(questions, board, class_level, subject, question_type, education_level="School"):
    groq_api_key = get_api_key("GROQ_API_KEY")
    if not groq_api_key:
        st.error("GROQ API key is missing. Please check your secrets.")
        return "Error: Missing API key. Contact the administrator."

    llm = ChatGroq(api_key=groq_api_key, model='llama-3.3-70b-versatile')

    answer_prompt = get_answer_prompt(question_type, board, class_level, subject, education_level)

    total_questions_tokens = estimate_tokens(questions)
    if total_questions_tokens > 3000:
        st.warning("Many questions detected. Generating answers in batches...")

        question_parts = re.split(r'(\d+\.\s)', questions)
        all_questions = []
        current_question = ""

        for part in question_parts:
            if re.match(r'\d+\.\s', part):
                if current_question:
                    all_questions.append(current_question)
                current_question = part
            else:
                current_question += part

        if current_question:
            all_questions.append(current_question)

        batch_size = max(1, len(all_questions) // 2)
        batches = [all_questions[i:i + batch_size] for i in range(0, len(all_questions), batch_size)]

        all_answers = []
        for i, batch in enumerate(batches):
            batch_text = "".join(batch)
            st.info(f"Generating answers for batch {i + 1} of {len(batches)}...")
            formatted_prompt = answer_prompt.format(questions=batch_text)
            try:
                response = llm.invoke(formatted_prompt)
                batch_answers = post_process_output(response.content)
                all_answers.append(batch_answers)
            except Exception as e:
                st.error(f"Error generating answers for batch {i + 1}: {e}")
                all_answers.append(
                    f"Error generating answers for questions "
                    f"{i * batch_size + 1}-{min((i + 1) * batch_size, len(all_questions))}"
                )

        combined_answers = "\n\n".join(all_answers)
        cleaned_answers = post_process_output(combined_answers)
    else:
        formatted_prompt = answer_prompt.format(questions=questions)
        with st.spinner(f'AI is generating answers for {question_type} questions...'):
            start_time = time.time()
            response = llm.invoke(formatted_prompt)
            generation_time = time.time() - start_time
            st.success(f"Generated answers in {generation_time:.1f} seconds!")

        cleaned_answers = post_process_output(response.content)

    st.session_state.generated_answers = cleaned_answers
    return cleaned_answers


def validate_questions(questions, board, class_level, subject, education_level="School"):
    groq_api_key = get_api_key("GROQ_API_KEY")
    if not groq_api_key:
        st.error("GROQ API key is missing. Please check your secrets.")
        return questions

    llm = ChatGroq(api_key=groq_api_key, model='llama-3.3-70b-versatile')
    level_label = get_level_label(class_level)

    if education_level == "College":
        curriculum_note = f"university/college {level_label} curriculum and academic standards"
    else:
        curriculum_note = f"{board} {level_label} curriculum"

    validation_prompt = f"""
    You are an expert validator for {curriculum_note} {subject} exam questions.

    Review the following questions and validate them based on these criteria:
    1. Relevance to {curriculum_note}
    2. Accuracy of content
    3. Clarity and unambiguous wording
    4. Appropriate difficulty level
    5. Grammatical correctness

    For each question, edit it directly if needed. Do not include reasoning or validation notes.
    Do not prefix the questions with "VALID:", "NEEDS REVISION:", or "REMOVE:".
    Simply output the corrected list of questions, starting with "1."

    Questions to validate:
    {questions}

    IMPORTANT FORMATTING INSTRUCTION: Output ONLY numbered questions starting with "1." - include ZERO preamble or explanations.
    Return the validated questions directly without any explanation of your validation process.
    """

    validation_tokens = estimate_tokens(validation_prompt)
    if validation_tokens > 4000:
        st.warning("Questions too long for full validation. Performing basic validation only.")
        max_question_tokens = 3000
        question_words = questions.split()
        truncated_questions = " ".join(question_words[:int(max_question_tokens / 1.3)])
        validation_prompt = validation_prompt.replace(questions, truncated_questions)

    try:
        with st.spinner('Validating questions...'):
            response = llm.invoke(validation_prompt)
        return response.content
    except Exception as e:
        st.error(f"Validation error: {e}")
        return questions


# ─── Main App ────────────────────────────────────────────────────────────────

def main():
    st.set_page_config(
        page_title='Question Generator - School & College',
        page_icon='📚',
        layout="wide"
    )

    init_session_state()

    st.title("Question Generator — School & College 📚")

    st.markdown("""
    <style>
    .main-header {color: #1E88E5; font-size: 28px; font-weight: bold;}
    .sub-header {color: #0277BD; font-size: 20px; font-weight: bold;}
    .highlight {background-color: #f0f7fb; padding: 10px; border-radius: 5px; border-left: 5px solid #2196F3;}
    .success {background-color: #e8f5e9; padding: 10px; border-radius: 5px; border-left: 5px solid #4CAF50;}
    .warning {background-color: #fff8e1; padding: 10px; border-radius: 5px; border-left: 5px solid #FF9800;}
    </style>
    """, unsafe_allow_html=True)

    groq_api_key = get_api_key("GROQ_API_KEY")
    if not groq_api_key:
        st.warning("GROQ API key is missing. The application may not function correctly.")
        st.info("For administrators: Please add GROQ_API_KEY to your Streamlit secrets.")

    # ── Sidebar ──────────────────────────────────────────────────────────────

    with st.sidebar:
        st.header("Upload PDF and Settings")

        # Education level toggle
        education_level = st.radio("🎓 Education Level:", EDUCATION_LEVELS, horizontal=True)
        st.session_state.education_level = education_level

        st.markdown("---")

        if education_level == "School":
            board = st.selectbox("Select Educational Board:", SCHOOL_BOARDS)
            class_level = st.selectbox("Select Class:", SCHOOL_CLASSES, index=9)
            subject = st.selectbox("Select Subject:", SCHOOL_SUBJECTS)
            question_types_list = SCHOOL_QUESTION_TYPES
            complexity_options = COMPLEXITY_LEVELS["School"]
        else:
            board = st.selectbox("Select Board / University Type:", COLLEGE_BOARDS)
            class_level = st.selectbox("Select Year:", COLLEGE_YEARS)
            subject = st.selectbox("Select Subject:", COLLEGE_SUBJECTS)
            question_types_list = COLLEGE_QUESTION_TYPES
            complexity_options = COMPLEXITY_LEVELS["College"]

        st.session_state.ocr_enabled = st.checkbox("Enable OCR for Scanned PDFs", value=True)

        st.markdown(
            '<p class="highlight">Select a PDF file containing curriculum content. '
            'Both text-based and scanned PDFs are supported.</p>',
            unsafe_allow_html=True
        )

        level_label = get_level_label(class_level)
        upload_pdf = st.file_uploader(
            f"Upload {board} {level_label} {subject} PDF",
            type="pdf",
            accept_multiple_files=False
        )

        st.markdown("### Customize Question Generation")
        question_types = st.selectbox("Select Question Type:", question_types_list, index=0)
        complexity_level = st.selectbox("Select Complexity Level:", complexity_options, index=1)
        num_questions = st.slider("Number of Questions:", min_value=5, max_value=25, value=10, step=1)
        validate = st.checkbox("Validate generated questions", value=True)

        st.markdown("### Performance Options")
        use_faster_model = st.checkbox("Use faster model (may reduce quality slightly)", value=True)

        # Question Bank
        if st.session_state.question_bank:
            st.markdown("### 📚 Question Bank")
            st.metric("Saved Questions", len(st.session_state.question_bank))
            with st.expander("View Saved"):
                for entry in st.session_state.question_bank[-5:]:
                    col_a, col_b = st.columns([3, 1])
                    with col_a:
                        st.caption(
                            f"{entry['settings']['subject']} "
                            f"({entry['settings'].get('education_level', 'School')}) - "
                            f"{entry['timestamp'][:10]}"
                        )
                    with col_b:
                        if st.button("📂", key=f"load_{entry['id']}"):
                            st.session_state.ai_answer = entry['questions']
                            st.session_state.answers = entry.get('answers', '')
                            st.session_state.answers_generated = bool(entry.get('answers'))
                            st.rerun()

        if upload_pdf:
            if st.button(f'Process {board} PDF', use_container_width=True):
                with st.spinner('Processing PDF...'):
                    st.info("Using HuggingFace embeddings")
                    st.session_state.embedding = load_embeddings()
                    st.session_state.raw_text = get_pdf_text(upload_pdf)

                    if st.session_state.raw_text:
                        st.markdown("### Preview of Extracted Text")
                        preview_text = (st.session_state.raw_text[:500] + "..."
                                        if len(st.session_state.raw_text) > 500
                                        else st.session_state.raw_text)
                        st.text_area("Extracted Text Preview", preview_text, height=100)

                        split_chunks = split_data_into_chunks(st.session_state.raw_text)
                        st.session_state.vectors = store_chunks_vectorDB(
                            chunks=split_chunks, embedding=st.session_state.embedding
                        )
                        st.success('PDF processing complete! 🎉')

                        word_count = len(st.session_state.raw_text.split())
                        st.markdown(
                            f"<p class='success'>Extracted {word_count} words from the PDF.</p>",
                            unsafe_allow_html=True
                        )

                        if 'answers' in st.session_state:
                            del st.session_state.answers
                        st.session_state.answers_generated = False
                    else:
                        st.error("Failed to extract text from the PDF. Please try another file.")

    # ── Main Content ─────────────────────────────────────────────────────────

    if st.session_state.vectors:
        st.subheader(f"Generate {board} {level_label} {subject} Questions ({education_level})")

        if hasattr(st.session_state, 'raw_text') and st.session_state.raw_text:
            word_count = len(st.session_state.raw_text.split())
            st.markdown(
                f"<p class='highlight'>Working with {word_count} words of content from your PDF.</p>",
                unsafe_allow_html=True
            )

        col1, col2 = st.columns(2)

        with col1:
            if st.button(f"Generate {question_types.title()} Questions", use_container_width=True):
                if 'answers' in st.session_state:
                    del st.session_state.answers
                st.session_state.answers_generated = False

                content = st.session_state.raw_text
                inputs = {
                    "content": content,
                    "question_types": question_types,
                    "complexity_level": complexity_level,
                    "num_questions": num_questions,
                    "board": board,
                    "class_level": class_level,
                    "subject": subject,
                    "education_level": education_level
                }

                st.session_state.ai_answer = generate_questions(inputs)

                if validate:
                    st.session_state.ai_answer = validate_questions(
                        st.session_state.ai_answer,
                        board, class_level, subject, education_level
                    )

                st.subheader(f"Generated {question_types.title()} Questions")
                st.markdown(
                    f"<div class='success'>{st.session_state.ai_answer.replace(chr(10), '<br>')}</div>",
                    unsafe_allow_html=True
                )

        with col2:
            if 'ai_answer' in st.session_state and not st.session_state.answers_generated:
                if st.button("Generate Answer Key", use_container_width=True):
                    with st.spinner("Generating detailed answers..."):
                        current_qt = question_types
                        if st.session_state.question_settings:
                            current_qt = st.session_state.question_settings.get("type", question_types)

                        st.session_state.answers = generate_answers(
                            st.session_state.ai_answer,
                            board, class_level, subject, current_qt, education_level
                        )
                        st.session_state.answers_generated = True

            if 'answers' in st.session_state and st.session_state.answers_generated:
                st.subheader("Answer Key")
                st.markdown(
                    f"<div class='success'>{st.session_state.answers.replace(chr(10), '<br>')}</div>",
                    unsafe_allow_html=True
                )

    # ── Question Editing ─────────────────────────────────────────────────────

    if 'ai_answer' in st.session_state and st.session_state.ai_answer:
        st.markdown("---")
        st.markdown("### ✏️ Edit Questions")

        edited_questions = st.text_area(
            "Review and edit your questions before exporting:",
            value=st.session_state.ai_answer,
            height=300,
            key="question_editor"
        )

        col_edit1, col_edit2, col_edit3 = st.columns(3)
        with col_edit1:
            if st.button("💾 Save Edits", use_container_width=True):
                st.session_state.ai_answer = edited_questions
                st.success("Questions updated!")

        with col_edit2:
            if st.button("🔄 Regenerate", use_container_width=True):
                st.rerun()

        with col_edit3:
            if st.button("📚 Save to Bank", use_container_width=True):
                if st.session_state.question_settings:
                    save_to_question_bank(
                        st.session_state.ai_answer,
                        st.session_state.get('answers', ''),
                        st.session_state.question_settings
                    )
                    st.success(f"Saved! Total in bank: {len(st.session_state.question_bank)}")

    # ── Export Options ───────────────────────────────────────────────────────

    if "ai_answer" in st.session_state and st.session_state.ai_answer:
        st.markdown("---")
        st.markdown("### 📥 Export Options")

        current_qt = question_types
        if st.session_state.question_settings:
            current_qt = st.session_state.question_settings.get("type", question_types)

        st.markdown("#### Export Questions")
        col1, col2, col3, col4, col5 = st.columns(5)

        with col1:
            if st.button("📄 PDF", use_container_width=True, key="export_q_pdf"):
                try:
                    with st.spinner("Creating PDF..."):
                        temp_pdf_path = create_temp_pdf(
                            st.session_state.ai_answer, board, class_level, subject, current_qt, is_answer_key=False
                        )
                        with open(temp_pdf_path, "rb") as pdf_file:
                            pdf_data = pdf_file.read()
                        os.unlink(temp_pdf_path)
                        st.download_button(
                            label="⬇️ Download PDF", data=pdf_data,
                            file_name=f"{board}_{level_label}_{subject}_{current_qt}_Questions.pdf",
                            mime="application/pdf", use_container_width=True, key="download_q_pdf"
                        )
                        st.success("PDF ready!")
                except Exception as e:
                    st.error(f"Error: {e}")

        with col2:
            if st.button("📝 Word", use_container_width=True, key="export_q_word"):
                try:
                    with st.spinner("Creating Word..."):
                        temp_docx_path = create_word_doc(
                            st.session_state.ai_answer, board, class_level, subject, current_qt, is_answer_key=False
                        )
                        with open(temp_docx_path, "rb") as docx_file:
                            docx_data = docx_file.read()
                        os.unlink(temp_docx_path)
                        st.download_button(
                            label="⬇️ Download Word", data=docx_data,
                            file_name=f"{board}_{level_label}_{subject}_{current_qt}_Questions.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True, key="download_q_word"
                        )
                        st.success("Word ready!")
                except Exception as e:
                    st.error(f"Error: {e}")

        with col3:
            if st.button("📊 Excel", use_container_width=True, key="export_q_excel"):
                try:
                    with st.spinner("Creating Excel..."):
                        temp_excel_path = create_excel_file(
                            st.session_state.ai_answer, board, class_level, subject, current_qt, is_answer_key=False
                        )
                        with open(temp_excel_path, "rb") as excel_file:
                            excel_data = excel_file.read()
                        os.unlink(temp_excel_path)
                        st.download_button(
                            label="⬇️ Download Excel", data=excel_data,
                            file_name=f"{board}_{level_label}_{subject}_{current_qt}_Questions.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            use_container_width=True, key="download_q_excel"
                        )
                        st.success("Excel ready!")
                except Exception as e:
                    st.error(f"Error: {e}")

        with col4:
            if st.button("📋 JSON", use_container_width=True, key="export_q_json"):
                try:
                    json_data = create_json_export(
                        st.session_state.ai_answer,
                        st.session_state.question_settings or {}
                    )
                    st.download_button(
                        label="⬇️ Download JSON", data=json_data,
                        file_name=f"{board}_{level_label}_{subject}_{current_qt}_Questions.json",
                        mime="application/json", use_container_width=True, key="download_q_json"
                    )
                    st.success("JSON ready!")
                except Exception as e:
                    st.error(f"Error: {e}")

        with col5:
            if st.button("📝 Text", use_container_width=True, key="export_q_text"):
                st.download_button(
                    label="⬇️ Download Text", data=st.session_state.ai_answer,
                    file_name=f"{board}_{level_label}_{subject}_{current_qt}_Questions.txt",
                    mime="text/plain", use_container_width=True, key="download_q_text"
                )
                st.success("Text ready!")

        # Export answers
        if 'answers' in st.session_state and st.session_state.answers_generated:
            st.markdown("#### Export Answer Key")
            col1, col2, col3, col4, col5 = st.columns(5)

            with col1:
                if st.button("📄 PDF", use_container_width=True, key="export_a_pdf"):
                    try:
                        with st.spinner("Creating PDF..."):
                            temp_pdf_path = create_temp_pdf(
                                st.session_state.answers, board, class_level, subject, current_qt, is_answer_key=True
                            )
                            with open(temp_pdf_path, "rb") as pdf_file:
                                pdf_data = pdf_file.read()
                            os.unlink(temp_pdf_path)
                            st.download_button(
                                label="⬇️ Download PDF", data=pdf_data,
                                file_name=f"{board}_{level_label}_{subject}_{current_qt}_AnswerKey.pdf",
                                mime="application/pdf", use_container_width=True, key="download_a_pdf"
                            )
                            st.success("PDF ready!")
                    except Exception as e:
                        st.error(f"Error: {e}")

            with col2:
                if st.button("📝 Word", use_container_width=True, key="export_a_word"):
                    try:
                        with st.spinner("Creating Word..."):
                            temp_docx_path = create_word_doc(
                                st.session_state.answers, board, class_level, subject, current_qt, is_answer_key=True
                            )
                            with open(temp_docx_path, "rb") as docx_file:
                                docx_data = docx_file.read()
                            os.unlink(temp_docx_path)
                            st.download_button(
                                label="⬇️ Download Word", data=docx_data,
                                file_name=f"{board}_{level_label}_{subject}_{current_qt}_AnswerKey.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True, key="download_a_word"
                            )
                            st.success("Word ready!")
                    except Exception as e:
                        st.error(f"Error: {e}")

            with col3:
                if st.button("📊 Excel", use_container_width=True, key="export_a_excel"):
                    try:
                        with st.spinner("Creating Excel..."):
                            temp_excel_path = create_excel_file(
                                st.session_state.answers, board, class_level, subject, current_qt, is_answer_key=True
                            )
                            with open(temp_excel_path, "rb") as excel_file:
                                excel_data = excel_file.read()
                            os.unlink(temp_excel_path)
                            st.download_button(
                                label="⬇️ Download Excel", data=excel_data,
                                file_name=f"{board}_{level_label}_{subject}_{current_qt}_AnswerKey.xlsx",
                                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                use_container_width=True, key="download_a_excel"
                            )
                            st.success("Excel ready!")
                    except Exception as e:
                        st.error(f"Error: {e}")

            with col4:
                if st.button("📋 JSON", use_container_width=True, key="export_a_json"):
                    try:
                        settings_copy = (st.session_state.question_settings.copy()
                                         if st.session_state.question_settings else {})
                        settings_copy['is_answer_key'] = True
                        json_data = create_json_export(st.session_state.answers, settings_copy)
                        st.download_button(
                            label="⬇️ Download JSON", data=json_data,
                            file_name=f"{board}_{level_label}_{subject}_{current_qt}_AnswerKey.json",
                            mime="application/json", use_container_width=True, key="download_a_json"
                        )
                        st.success("JSON ready!")
                    except Exception as e:
                        st.error(f"Error: {e}")

            with col5:
                if st.button("📝 Text", use_container_width=True, key="export_a_text"):
                    st.download_button(
                        label="⬇️ Download Text", data=st.session_state.answers,
                        file_name=f"{board}_{level_label}_{subject}_{current_qt}_AnswerKey.txt",
                        mime="text/plain", use_container_width=True, key="download_a_text"
                    )
                    st.success("Text ready!")

    # ── Getting Started ──────────────────────────────────────────────────────

    if not st.session_state.vectors:
        st.markdown("## 🚀 How to Get Started")
        st.markdown("""
        1. **Choose Education Level**: Select **School** (Classes 1-12) or **College** (1st Year – Post Graduate)
        2. **Upload a PDF**: Select a PDF containing curriculum material for your chosen subject
        3. **Click 'Process PDF'**: The system will extract text (supports both text PDFs and scanned/image PDFs)
        4. **Configure Questions**: Choose the question type, complexity level, and number of questions
        5. **Generate Questions**: Click the button to create board/university-specific questions
        6. **Create Answer Key**: Generate detailed answers for the questions
        7. **Export**: Download in multiple formats — PDF, Word, Excel, JSON, or Text
        8. **Save to Bank**: Store frequently used questions for future reference
        """)

        col_info1, col_info2 = st.columns(2)
        with col_info1:
            st.info(
                "🏫 **School Mode** supports CBSE, ICSE & State Board for Classes 1-12 "
                "with standard question types."
            )
        with col_info2:
            st.info(
                "🎓 **College Mode** supports University/UGC/AICTE/Autonomous boards with "
                "advanced question types like Case Studies, Numerical Problems, "
                "Assertion-Reason, and Conceptual Analysis."
            )

        st.info(
            "This tool supports OCR for scanned PDFs and handwritten content. "
            "Enable the OCR option in the sidebar if you're using scanned materials."
        )


if __name__ == "__main__":
    main()
